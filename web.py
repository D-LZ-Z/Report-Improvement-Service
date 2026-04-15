import math
import os
import time
from io import BytesIO

import streamlit as st
from dotenv import load_dotenv
from docx import Document
import PyPDF2
from openai import OpenAI
from playwright.sync_api import sync_playwright

# =========================
# Stage 3A: Setup
# 负责：初始化页面和环境变量
# =========================
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
STEALTHWRITER_EMAIL = os.getenv("STEALTHWRITER_EMAIL", "")
STEALTHWRITER_PASSWORD = os.getenv("STEALTHWRITER_PASSWORD", "")

client = OpenAI(api_key=OPENAI_API_KEY)

st.set_page_config(
    page_title="Assignment Improvement Service",
    page_icon="📝",
    layout="wide"
)

# =========================
# Stage 3B: Styling
# 负责：页面视觉和升级转化设计
# =========================
st.markdown("""
<style>
.block-container {
    max-width: 1120px;
    padding-top: 2rem;
    padding-bottom: 3rem;
}
.hero-box {
    background: linear-gradient(135deg, #111827, #374151);
    color: white;
    padding: 28px;
    border-radius: 22px;
    margin-bottom: 24px;
}
.hero-title {
    font-size: 2rem;
    font-weight: 700;
    margin-bottom: 6px;
}
.hero-subtitle {
    color: #e5e7eb;
}
.word-box, .info-box, .success-box, .warning-box, .preview-box {
    border-radius: 16px;
    padding: 16px;
    margin-top: 12px;
    margin-bottom: 16px;
}
.word-box, .info-box, .preview-box {
    background: #f9fafb;
    border: 1px solid #e5e7eb;
}
.success-box {
    background: #ecfdf5;
    border: 1px solid #a7f3d0;
    color: #065f46;
}
.warning-box {
    background: #fff7ed;
    border: 1px solid #fdba74;
    color: #9a3412;
}
.plan-card {
    border: 1.5px solid #e5e7eb;
    border-radius: 18px;
    padding: 20px;
    background: white;
    min-height: 340px;
}
.plan-card-popular {
    border: 2px solid #111827;
    border-radius: 18px;
    padding: 20px;
    background: #f9fafb;
    min-height: 340px;
    box-shadow: 0 10px 24px rgba(0,0,0,0.07);
}
.badge {
    display: inline-block;
    background: #111827;
    color: white;
    border-radius: 999px;
    padding: 4px 10px;
    font-size: 0.8rem;
    font-weight: 600;
    margin-bottom: 10px;
}
.plan-title {
    font-size: 1.25rem;
    font-weight: 700;
}
.plan-price {
    font-size: 2rem;
    font-weight: 800;
    margin: 10px 0;
}
.plan-desc {
    color: #6b7280;
    margin-bottom: 14px;
}
.small-note {
    color: #6b7280;
    font-size: 0.92rem;
}
</style>
""", unsafe_allow_html=True)

st.markdown("""
<div class="hero-box">
  <div class="hero-title">Assignment Improvement Service</div>
  <div class="hero-subtitle">
    Upload your assignment, preview the first 350 words, choose your service, and receive your improved version after payment.
  </div>
</div>
""", unsafe_allow_html=True)

# =========================
# Stage 3C: Helper Functions
# 负责：读文件、算字数、价格、导出
# =========================
def read_uploaded_file(file) -> str:
    if file.type == "text/plain":
        return file.read().decode("utf-8", errors="ignore")

    if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs if p.text])

    if file.type == "application/pdf":
        pdf = PyPDF2.PdfReader(file)
        pages = []
        for page in pdf.pages:
            content = page.extract_text()
            if content:
                pages.append(content)
        return "\n".join(pages)

    return ""


def count_words(text: str) -> int:
    return len(text.split())


def calculate_price(words: int, plan_key: str) -> int:
    if plan_key == "content":
        if words <= 1000:
            return 15
        elif words <= 2000:
            return 25
        elif words <= 3000:
            return 35
        elif words <= 12000:
            extra_blocks = math.ceil((words - 3000) / 1000)
            return 35 + (extra_blocks * 10)
        elif words <= 25000:
            return 100
        else:
            extra_blocks = math.ceil((words - 25000) / 1000)
            return 100 + (extra_blocks * 8)

    if plan_key == "score_boost":
        if words <= 1000:
            return 20
        elif words <= 2000:
            return 30
        elif words <= 3000:
            return 45
        elif words <= 12000:
            extra_blocks = math.ceil((words - 3000) / 1000)
            return 45 + (extra_blocks * 12)
        elif words <= 25000:
            return 120
        else:
            extra_blocks = math.ceil((words - 25000) / 1000)
            return 120 + (extra_blocks * 10)

    if plan_key == "humanised":
        if words <= 1000:
            return 25
        elif words <= 2000:
            return 40
        elif words <= 3000:
            return 50
        elif words <= 12000:
            extra_blocks = math.ceil((words - 3000) / 1000)
            return 50 + (extra_blocks * 15)
        elif words <= 25000:
            return 150
        else:
            extra_blocks = math.ceil((words - 25000) / 1000)
            return 150 + (extra_blocks * 12)

    return 0


def build_preview(text: str, limit_words: int = 350) -> str:
    return " ".join(text.split()[:limit_words])


def export_txt(text: str) -> bytes:
    return text.encode("utf-8")


def export_docx(text: str) -> BytesIO:
    output = BytesIO()
    doc = Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output)
    output.seek(0)
    return output

# =========================
# Stage 3D: Improvement Engine
# 负责：内容优化
# =========================
def improve_text(text: str, level: str = "content") -> str:
    if level == "content":
        system_prompt = (
            "You are a professional academic editor. Improve the user's writing by fixing grammar, "
            "improving clarity, and making sentence flow cleaner while preserving meaning. "
            "Do not invent facts, citations, or references."
        )
    elif level == "score_boost":
        system_prompt = (
            "You are a professional academic editor. Improve the user's writing more deeply by fixing grammar, "
            "strengthening clarity, improving wording, transitions, paragraph flow, and overall academic readability. "
            "Do not invent facts, citations, or references."
        )
    else:
        system_prompt = (
            "You are a professional academic editor. Improve the user's writing comprehensively, making it smoother, "
            "more polished, more natural, and more refined while preserving the original meaning. "
            "Do not invent facts, citations, or references."
        )

    response = client.chat.completions.create(
        model="gpt-4o-mini",
        temperature=0.4,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": text},
        ],
    )
    return response.choices[0].message.content.strip()

# =========================
# Stage 3E: Humanised Step
# 负责：高级方案接 StealthWriter
# =========================
def humanise_with_stealthwriter(text: str) -> str:
    if not STEALTHWRITER_EMAIL or not STEALTHWRITER_PASSWORD:
        raise RuntimeError("Missing StealthWriter credentials in .env")

    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        page.goto("https://stealthwriter.ai/", wait_until="domcontentloaded")
        time.sleep(3)

        # 下面这些 selector 需要你自己 inspect 后改
        # page.click("text=Login")
        # page.fill("input[type='email']", STEALTHWRITER_EMAIL)
        # page.fill("input[type='password']", STEALTHWRITER_PASSWORD)
        # page.click("button[type='submit']")
        # time.sleep(5)

        page.fill("textarea", text)
        time.sleep(1)
        page.click("button")
        time.sleep(8)

        result = page.locator("textarea").nth(1).input_value()
        browser.close()
        return result

# =========================
# Stage 3F: Session State
# 负责：保存流程状态
# =========================
defaults = {
    "raw_text": "",
    "word_count": 0,
    "preview_text": "",
    "selected_plan": "",
    "selected_price": 0,
    "payment_uploaded": False,
    "final_result": "",
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

# =========================
# Stage 3G: Upload Assignment
# 负责：上传文件
# =========================
st.subheader("1. Upload Your Assignment")

uploaded_file = st.file_uploader(
    "Supported formats: .txt, .docx, .pdf",
    type=["txt", "docx", "pdf"]
)

if uploaded_file:
    text = read_uploaded_file(uploaded_file).strip()
    st.session_state["raw_text"] = text
    st.session_state["word_count"] = count_words(text)

    st.markdown(
        f'<div class="word-box"><b>Word Count:</b> {st.session_state["word_count"]} words</div>',
        unsafe_allow_html=True
    )

    if st.session_state["word_count"] < 50:
        st.markdown(
            '<div class="warning-box">The uploaded file is too short. Please upload a longer assignment.</div>',
            unsafe_allow_html=True
        )

# =========================
# Stage 3H: Preview
# 负责：显示 350 字 preview
# =========================
if st.session_state["raw_text"] and st.session_state["word_count"] >= 50:
    st.subheader("2. Preview")

    if st.button("Generate Preview"):
        with st.spinner("Preparing preview..."):
            preview_version = improve_text(st.session_state["raw_text"], level="content")
            st.session_state["preview_text"] = build_preview(preview_version, 350)

    if st.session_state["preview_text"]:
        st.markdown('<div class="preview-box">', unsafe_allow_html=True)
        st.write(st.session_state["preview_text"])
        st.markdown('</div>', unsafe_allow_html=True)
        st.caption("Preview shows the first 350 words only.")

# =========================
# Stage 3I: Pricing UI
# 负责：学生消费习惯升级设计
# =========================
if st.session_state["preview_text"]:
    st.subheader("3. Choose Your Service")

    wc = st.session_state["word_count"]
    price_content = calculate_price(wc, "content")
    price_score = calculate_price(wc, "score_boost")
    price_humanised = calculate_price(wc, "humanised")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown(f"""
        <div class="plan-card">
            <div class="plan-title">Content Improvements</div>
            <div class="plan-price">RM{price_content}</div>
            <div class="plan-desc">Best for simple cleanup</div>
            <div>• Improve clarity</div>
            <div>• Fix grammar issues</div>
            <div>• Better sentence flow</div>
            <div>• Basic polishing for submission</div>
            <br>
            <div class="small-note">Suitable if you only want a lighter improvement.</div>
        </div>
        """, unsafe_allow_html=True)

        if st.button("Choose Content Improvements", use_container_width=True):
            st.session_state["selected_plan"] = "content"
            st.session_state["selected_price"] = price_content

    with col2:
        st.markdown(f"""
        <div class="plan-card-popular">
            <div class="badge">Most Popular</div>
            <div class="plan-title">Score Boost</div>
            <div class="plan-price">RM{price_score}</div>
            <div class="plan-desc">Best value for students</div>
            <div>• Stronger wording</div>
            <div>• Better readability and structure</div>
            <div>• Cleaner academic flow</div>
            <div>• More polished submission quality</div>
            <br>
            <div class="small-note">Students usually choose this for a stronger result without paying too much.</div>
        </div>
        """, unsafe_allow_html=True)

        if st.button("Choose Score Boost", use_container_width=True):
            st.session_state["selected_plan"] = "score_boost"
            st.session_state["selected_price"] = price_score

    with col3:
        st.markdown(f"""
        <div class="plan-card">
            <div class="plan-title">Improvements + Humanised</div>
            <div class="plan-price">RM{price_humanised}</div>
            <div class="plan-desc">Best for the strongest refinement</div>
            <div>• Deeper rewriting</div>
            <div>• More natural expression</div>
            <div>• Smoother and more polished writing</div>
            <div>• Highest level of refinement</div>
            <br>
            <div class="small-note">Recommended if you want the strongest final refinement.</div>
        </div>
        """, unsafe_allow_html=True)

        if st.button("Choose Improvements + Humanised", use_container_width=True):
            st.session_state["selected_plan"] = "humanised"
            st.session_state["selected_price"] = price_humanised

    if st.session_state["selected_plan"]:
        selected_name = {
            "content": "Content Improvements",
            "score_boost": "Score Boost",
            "humanised": "Improvements + Humanised"
        }[st.session_state["selected_plan"]]

        st.markdown(
            f'<div class="success-box"><b>Selected Service:</b> {selected_name}<br><b>Total Price:</b> RM{st.session_state["selected_price"]}</div>',
            unsafe_allow_html=True
        )

# =========================
# Stage 3J: Payment
# 负责：显示二维码 + 上传付款证明
# =========================
if st.session_state["selected_plan"]:
    st.subheader("4. Payment")

    st.write(f"Please pay **RM{st.session_state['selected_price']}** using one of the QR codes below.")

    q1, q2 = st.columns(2)

    with q1:
        st.markdown("**Touch 'n Go**")
        st.image("tng_qr.png", use_container_width=True)

    with q2:
        st.markdown("**Bank Transfer / DuitNow**")
        st.image("bank_qr.png", use_container_width=True)

    payment_proof = st.file_uploader(
        "Upload payment proof",
        type=["png", "jpg", "jpeg", "pdf"],
        key="payment_proof"
    )

    if payment_proof:
        st.session_state["payment_uploaded"] = True
        st.markdown(
            '<div class="success-box">Payment proof uploaded. You can now generate the final result.</div>',
            unsafe_allow_html=True
        )

# =========================
# Stage 3K: Final Output
# 负责：付款后生成完整结果
# =========================
if st.session_state["payment_uploaded"] and st.session_state["selected_plan"]:
    st.subheader("5. Generate Final Result")

    if st.button("Generate Final Result"):
        with st.spinner("Processing your document..."):
            source_text = st.session_state["raw_text"]
            plan = st.session_state["selected_plan"]

            if plan == "content":
                final_result = improve_text(source_text, level="content")
            elif plan == "score_boost":
                final_result = improve_text(source_text, level="score_boost")
            else:
                improved = improve_text(source_text, level="humanised")
                try:
                    final_result = humanise_with_stealthwriter(improved)
                except Exception as e:
                    st.warning(f"Automatic humanised step could not complete: {e}")
                    st.info("Returning the improved version instead.")
                    final_result = improved

            st.session_state["final_result"] = final_result

    if st.session_state["final_result"]:
        st.subheader("6. Final Result")
        st.text_area("Your improved content", st.session_state["final_result"], height=420)

        txt_file = export_txt(st.session_state["final_result"])
        docx_file = export_docx(st.session_state["final_result"])

        d1, d2 = st.columns(2)

        with d1:
            st.download_button(
                label="Download as .txt",
                data=txt_file,
                file_name="improved_assignment.txt",
                mime="text/plain",
                use_container_width=True
            )

        with d2:
            st.download_button(
                label="Download as .docx",
                data=docx_file,
                file_name="improved_assignment.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )